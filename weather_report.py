from docx import Document
from docx.shared import Inches

document = Document()
document.add_heading('Weather in Famous Cities', 0)

p = document.add_paragraph('This projects aims to demonstrate automatically create ')
p.add_run('Word').bold = True
p.add_run(' file with ')
p.add_run('Python.').italic = True

#weather data from  https://www.timeanddate.com/weather/?continent=asia&sort=1
# city, weather, temperature °C
records = (
    ('Hong Kong', 'Sunny', 29),
    ('New Work', 'Clear', 0),
    ('London', 'Overcast', 7),
    ('Tokyo', 'Scattered clouds', 19),
    ('Taipei', 'Scattered clouds', 27),
    ('Ho Chi Minh', 'Scattered clouds', 34),
)


table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'City'
hdr_cells[1].text = 'Weather'
hdr_cells[2].text = 'Temperature °C'

for city, weather, temp in records:
    row_cells = table.add_row().cells
    row_cells[0].text = city
    row_cells[1].text = weather
    row_cells[2].text = str(temp)


document.add_picture('weather.png', width=Inches(5))


# https://python-docx.readthedocs.io/en/latest/user/hdrftr.html
section = document.sections[0]
header = section.header
paragraph = header.paragraphs[0]
paragraph.text = "Created by Roger Lau"

document.save('weather in cities.docx')